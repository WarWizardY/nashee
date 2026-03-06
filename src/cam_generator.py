from pathlib import Path
from typing import Dict, Any

from docx import Document
from docx.shared import Inches

from .risk_engine import RiskDecision


def generate_cam_docx(
    output_path: Path,
    company_name: str,
    sector: str,
    requested_limit: float,
    risk_decision: RiskDecision,
    input_summary: Dict[str, Any],
    application_id: str | None = None,
    engine_version: str | None = None,
) -> None:
    """
    Generate a simple but structured Credit Appraisal Memo (CAM) in DOCX format.
    This is a prototype aligned with the Five Cs of Credit.
    """
    doc = Document()

    # Title
    doc.add_heading(f"Credit Appraisal Memo – {company_name}", level=1)

    # Meta / lineage
    meta_para = doc.add_paragraph()
    meta_para.add_run("Application ID: ").bold = True
    meta_para.add_run(application_id or "N/A")
    meta_para.add_run("   Engine Version: ").bold = True
    meta_para.add_run(engine_version or "N/A")

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=2)
    decision_text = "APPROVE" if risk_decision.approve else "REJECT"
    p = doc.add_paragraph()
    p.add_run("Decision: ").bold = True
    p.add_run(decision_text)

    p = doc.add_paragraph()
    p.add_run("Requested Limit: ").bold = True
    p.add_run(f"INR {requested_limit:,.0f}")

    p = doc.add_paragraph()
    p.add_run("Recommended Limit: ").bold = True
    p.add_run(f"INR {risk_decision.recommended_limit:,.0f}")

    p = doc.add_paragraph()
    p.add_run("Recommended Interest Rate: ").bold = True
    if risk_decision.approve:
        p.add_run(f"{risk_decision.recommended_rate:.2f}%")
    else:
        p.add_run("N/A (rejected)")

    p = doc.add_paragraph()
    p.add_run("Overall Risk Score: ").bold = True
    p.add_run(f"{risk_decision.score:.2f} (0–1 scale)")
    p = doc.add_paragraph()
    p.add_run("Risk Band: ").bold = True
    p.add_run(risk_decision.risk_band)
    p = doc.add_paragraph()
    p.add_run("PD Estimate (prototype): ").bold = True
    p.add_run(f"{risk_decision.pd_estimate:.2%}")

    # Basic company info
    doc.add_heading("2. Company & Sector Overview", level=2)
    p = doc.add_paragraph()
    p.add_run("Company: ").bold = True
    p.add_run(company_name)
    p = doc.add_paragraph()
    p.add_run("Sector: ").bold = True
    p.add_run(sector)

    # Five Cs of Credit (prototype)
    doc.add_heading("3. Five Cs of Credit Assessment", level=2)

    # Character
    doc.add_heading("3.1 Character", level=3)
    
    # 3.1.1 Research Agent & NLP Signals
    doc.add_heading("Research Agent & NLP Signals", level=4)
    res_table = doc.add_table(rows=1, cols=2)
    rhdr = res_table.rows[0].cells
    rhdr[0].text = "Signal"
    rhdr[1].text = "Observation"
    
    def _r_add_row(metric: str, value: str) -> None:
        row_cells = res_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value

    if "news_sentiment_score" in input_summary:
        _r_add_row("News Sentiment Score", f"{input_summary['news_sentiment_score']:.2f} (from FinBERT)")
    if "research_mca_status" in input_summary:
        _r_add_row("MCA Company Status", str(input_summary["research_mca_status"]))
    if "research_ecourts_litigation_count" in input_summary:
        _r_add_row("e-Courts Litigation Count", str(input_summary["research_ecourts_litigation_count"]))
    
    doc.add_paragraph(
        "Note: Character assessment incorporates the above automated web-scraping "
        "and FinBERT NLP sentiment signals."
    )

    # Capacity
    doc.add_heading("3.2 Capacity", level=3)
    cap_para = doc.add_paragraph()
    cap_para.add_run("Financial Capacity Snapshot (latest year):").bold = True

    fin_table = doc.add_table(rows=1, cols=2)
    hdr_cells = fin_table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Value"

    def _add_row(metric: str, value: str) -> None:
        row_cells = fin_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value

    _add_row("Revenue", f"INR {input_summary.get('latest_revenue', 0.0):,.0f}")
    _add_row("EBITDA", f"INR {input_summary.get('latest_ebitda', 0.0):,.0f}")
    _add_row("PAT", f"INR {input_summary.get('latest_pat', 0.0):,.0f}")
    _add_row("Net Worth", f"INR {input_summary.get('latest_net_worth', 0.0):,.0f}")
    _add_row("Total Debt", f"INR {input_summary.get('latest_total_debt', 0.0):,.0f}")

    # If we have sanction letter derived features, surface them explicitly
    sanction_amount = input_summary.get("sanction_loan_amount")
    if sanction_amount is not None:
        doc.add_paragraph().add_run("Existing Sanctioned Facility (from sanction letter):").bold = True
        sanction_table = doc.add_table(rows=1, cols=2)
        shdr = sanction_table.rows[0].cells
        shdr[0].text = "Attribute"
        shdr[1].text = "Value"

        def _s_add_row(metric: str, value: str) -> None:
            row_cells = sanction_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value

        _s_add_row("Sanctioned Amount", f"INR {sanction_amount:,.0f}")
        if input_summary.get("sanction_interest_rate") is not None:
            _s_add_row("Interest Rate", f"{input_summary['sanction_interest_rate']:.2f}%")
        if input_summary.get("sanction_tenure_months") is not None:
            _s_add_row("Tenure", f"{int(input_summary['sanction_tenure_months'])} months")
        if input_summary.get("sanction_facility_type"):
            _s_add_row("Facility Type", str(input_summary["sanction_facility_type"]))
        if input_summary.get("sanction_guarantee_type"):
            _s_add_row("Guarantee Type", str(input_summary["sanction_guarantee_type"]))
        if input_summary.get("sanction_bank_name"):
            _s_add_row("Bank Name", str(input_summary["sanction_bank_name"]))

    doc.add_paragraph(
        "Note: In this prototype, capacity assessment is based on simple ratios. "
        "Full implementation will add GST/bank cash-flow analysis and DSCR metrics."
    )

    # Capital
    doc.add_heading("3.3 Capital", level=3)
    doc.add_paragraph(
        "Prototype: Capital strength is approximated via Net Worth and leverage. "
        "Future versions will incorporate detailed balance sheet trends and retained earnings."
    )

    # Collateral
    doc.add_heading("3.4 Collateral", level=3)
    doc.add_paragraph(
        "Prototype: Collateral information is not yet integrated. "
        "Full version will parse sanction letters, security documents, and MCA charges."
    )

    # Conditions
    doc.add_heading("3.5 Conditions", level=3)
    cond_para = doc.add_paragraph()
    cond_para.add_run(
        "Prototype: Sectoral and regulatory conditions are not yet fully modeled. "
        "Future versions will leverage RBI/NBFC guidelines, sector research, and external news."
    )

    # If transaction graph evidence is available, surface it here
    graph_cycles = input_summary.get("graph_example_cycles") or []
    top_central = input_summary.get("graph_top_central_entities") or []
    if graph_cycles or top_central:
        doc.add_paragraph("Transaction network observations:", style="List Bullet")
        if graph_cycles:
            for cycle in graph_cycles:
                doc.add_paragraph(
                    f"Potential circular trading loop detected among entities: {' → '.join(map(str, cycle))}.",
                    style="List Bullet 2",
                )
        if top_central:
            for entity, centrality in top_central:
                doc.add_paragraph(
                    f"Entity {entity} shows high transactional centrality (in-degree centrality={centrality:.2f}).",
                    style="List Bullet 2",
                )

    # Model Rationale
    doc.add_heading("4. Model Rationale & Key Factors", level=2)
    doc.add_paragraph(
        "Below are the main factors from the prototype rule-based engine that influenced the decision:"
    )
    for reason in risk_decision.reasons:
        doc.add_paragraph(reason, style="List Bullet")

    # Data Sources & Confidence
    doc.add_heading("5. Data Sources & Confidence (Prototype)", level=2)
    sources = input_summary.get("data_sources_present", [])
    completeness = input_summary.get("data_completeness_score", 0.0)
    p = doc.add_paragraph()
    p.add_run("Data Sources Provided: ").bold = True
    p.add_run(", ".join(sources) if sources else "None")
    p = doc.add_paragraph()
    p.add_run("Data Completeness Score: ").bold = True
    p.add_run(f"{completeness:.2f} (0–1, higher = more inputs available)")

    # Optional graph image
    graph_image_path = input_summary.get("graph_image_path")
    if graph_image_path:
        doc.add_heading("6. Transaction Graph (Prototype Visualization)", level=2)
        try:
            doc.add_picture(graph_image_path, width=Inches(4.5))
        except Exception:
            doc.add_paragraph(f"(Graph image could not be loaded from {graph_image_path})")

    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

